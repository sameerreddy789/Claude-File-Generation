"""
Fix the 2 remaining corrections by working directly on runs.
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document

doc = Document("8th Social Text Book SEm1 2026_CORRECTED.docx")

# FIX 1: PARA 7841 - "Prataparudra- II \u2018s time" → "Prataparudra II\u2019s time"
para = doc.paragraphs[7841]
run3 = para.runs[3]
print(f"7841 Run 3 before: '{run3.text}'")
# The text is "Prataparudra- II \u2018s time"
# Fix: remove the dash and extra space, use proper apostrophe
run3.text = "Prataparudra II\u2019s time"
print(f"7841 Run 3 after:  '{run3.text}'")

# FIX 2: PARA 7903 - "explains about Kakatiya\u2019s" → "explain Kakatiya\u2019s"
para = doc.paragraphs[7903]
# Run 10: "explains" → "explain"
print(f"\n7903 Run 10 before: '{para.runs[10].text}'")
para.runs[10].text = "explain"
print(f"7903 Run 10 after:  '{para.runs[10].text}'")

# Run 12: "about" → "" (remove)
print(f"7903 Run 12 before: '{para.runs[12].text}'")
para.runs[12].text = ""
print(f"7903 Run 12 after:  '{para.runs[12].text}'")

# Run 13: " " → "" (remove extra space after 'about')
print(f"7903 Run 13 before: '{para.runs[13].text}'")
para.runs[13].text = ""
print(f"7903 Run 13 after:  '{para.runs[13].text}'")

# Verify final text
full_7903 = ''.join(r.text for r in doc.paragraphs[7903].runs)
print(f"\n7903 final: {full_7903}")

full_7841 = ''.join(r.text for r in doc.paragraphs[7841].runs)
print(f"7841 final: {full_7841[:120]}")

# Save
output = "8th Social Text Book SEm1 2026_CORRECTED.docx"
doc.save(output)
print(f"\nSaved: {output}")
print("All 51 corrections now applied!")
