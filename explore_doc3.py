"""
Focused exploration: skip headers/footers, just search paragraphs.
Find content between printed page numbers 222-242.
"""
from docx import Document
import re
import sys

sys.setrecursionlimit(5000)

doc = Document("8th Social Text Book SEm1 2026.docx")

print(f"Total paragraphs: {len(doc.paragraphs)}")

# 1. Search ALL paragraphs for text containing "222"
print("\n=== PARAGRAPHS CONTAINING '222' ===")
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if '222' in text:
        has_telugu = bool(re.search(r'[\u0C00-\u0C7F]', text))
        lang = "TE" if has_telugu else "EN"
        print(f"  [{i}] [{lang}] Style='{para.style.name}' | '{text[:120]}'")

print("\n=== PARAGRAPHS CONTAINING '242' ===")
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if '242' in text:
        has_telugu = bool(re.search(r'[\u0C00-\u0C7F]', text))
        lang = "TE" if has_telugu else "EN"
        print(f"  [{i}] [{lang}] Style='{para.style.name}' | '{text[:120]}'")

# 2. Content from paragraph 7014 (page 207) onward - show 120 paras
print("\n=== CONTENT FROM PARA 7014 (page 207) - next 120 paras with text ===")
shown = 0
for i in range(7014, min(len(doc.paragraphs), 9000)):
    para = doc.paragraphs[i]
    text = para.text.strip()
    if text and shown < 120:
        has_telugu = bool(re.search(r'[\u0C00-\u0C7F]', text))
        lang = "TE" if has_telugu else "EN"
        display = text[:120] + ("..." if len(text) > 120 else "")
        print(f"  [{i}] [{lang}] Style='{para.style.name}' | {display}")
        shown += 1

print("\nDone!")
