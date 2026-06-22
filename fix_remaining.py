"""
Fix the 2 remaining corrections that failed due to run-level text issues.
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document

doc = Document("8th Social Text Book SEm1 2026_CORRECTED.docx")

# Debug: show runs for paragraph 7841 and 7903
print("=== PARA 7841 RUNS ===")
para = doc.paragraphs[7841]
for i, run in enumerate(para.runs):
    print(f"  Run {i}: '{run.text}' | chars: {[hex(ord(c)) for c in run.text[:30]]}")

print("\n=== PARA 7903 RUNS ===")
para = doc.paragraphs[7903]
for i, run in enumerate(para.runs):
    print(f"  Run {i}: '{run.text}' | chars: {[hex(ord(c)) for c in run.text[:30]]}")

# Now fix them by working directly on runs
# PARA 7841: Fix "Prataparudra- II 's" spacing
para = doc.paragraphs[7841]
full = ''.join(r.text for r in para.runs)
print(f"\n7841 full text: {full[:100]}")
print(f"7841 chars: {[hex(ord(c)) for c in full[:40]]}")

# PARA 7903: Fix "explains about Kakatiya's"
para = doc.paragraphs[7903]
full = ''.join(r.text for r in para.runs)
print(f"\n7903 full text: {full[:150]}")
print(f"7903 chars at 'explains': ", end='')
idx = full.find('explains')
if idx >= 0:
    print([hex(ord(c)) for c in full[idx:idx+40]])
else:
    print("NOT FOUND in runs")
    print(f"para.text: {para.text[:150]}")
    print(f"para.text chars at 'explains': ", end='')
    idx2 = para.text.find('explains')
    if idx2 >= 0:
        print([hex(ord(c)) for c in para.text[idx2:idx2+40]])
