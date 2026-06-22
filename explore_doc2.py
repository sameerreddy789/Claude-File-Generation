"""
Deeper exploration: find page numbers 222-242.
Check headers, footers, sections, and different text patterns.
"""
from docx import Document
from docx.oxml.ns import qn
import re

doc = Document("8th Social Text Book SEm1 2026.docx")

# 1. Check sections and their headers/footers
print("=== SECTIONS ===")
for i, section in enumerate(doc.sections):
    print(f"Section {i}: start_type={section.start_type}")
    # Check headers
    if section.header:
        for para in section.header.paragraphs:
            if para.text.strip():
                print(f"  Header: '{para.text.strip()[:80]}'")
    if section.footer:
        for para in section.footer.paragraphs:
            if para.text.strip():
                print(f"  Footer: '{para.text.strip()[:80]}'")

# 2. Look for page numbers in different formats
print("\n=== SEARCHING FOR '222' and '242' IN ALL PARAGRAPHS ===")
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if '222' in text and len(text) < 200:
        print(f"  [{i}] '{text[:150]}' | Style='{para.style.name}'")

print("\n--- Looking for 242 ---")
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if '242' in text and len(text) < 200:
        print(f"  [{i}] '{text[:150]}' | Style='{para.style.name}'")

# 3. Search for page break elements in XML 
print("\n=== PAGE BREAKS (first 30) ===")
body = doc.element.body
page_breaks = body.findall('.//' + qn('w:br'))
count = 0
for br in page_breaks:
    br_type = br.get(qn('w:type'))
    if br_type == 'page':
        # Find parent paragraph
        parent = br.getparent()
        while parent is not None and parent.tag != qn('w:p'):
            parent = parent.getparent()
        if parent is not None:
            # Find index
            for idx, para in enumerate(doc.paragraphs):
                if para._element == parent:
                    print(f"  Page break at paragraph {idx}")
                    count += 1
                    break
        if count >= 30:
            break

# 4. Look at content around paragraph 7014 (page 207) and forward
# Since page 207 is at index 7014, page 222 would be roughly 15 pages later
# Let's look at content from 7014 onwards to find the pattern
print("\n=== CONTENT FROM PARAGRAPH 7014 ONWARDS (first 80 paras with text) ===")
shown = 0
for i in range(7014, min(len(doc.paragraphs), 8500)):
    para = doc.paragraphs[i]
    text = para.text.strip()
    if text and shown < 80:
        # Detect if text contains Telugu (Unicode range for Telugu: 0C00-0C7F)
        has_telugu = bool(re.search(r'[\u0C00-\u0C7F]', text))
        lang = "TELUGU" if has_telugu else "ENGLISH"
        display = text[:100] + ("..." if len(text) > 100 else "")
        print(f"  [{i}] [{lang}] Style='{para.style.name}' | {display}")
        shown += 1
