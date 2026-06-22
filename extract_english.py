"""
Extract ALL English text paragraphs between paragraphs 7300-8172
(approximate range for printed pages 222-242).
Show full text for grammar correction.
"""
from docx import Document
import re

doc = Document("8th Social Text Book SEm1 2026.docx")

def is_real_english(text):
    """Check if text is actual English content (not dots/spaces pattern or Telugu)."""
    stripped = text.strip()
    if not stripped:
        return False
    # Has Telugu characters?
    if re.search(r'[\u0C00-\u0C7F]', stripped):
        return False
    # Is it just dots, spaces, tabs?
    cleaned = re.sub(r'[.\s\t\r\n]', '', stripped)
    if len(cleaned) < 3:
        return False
    # Has actual English letters?
    if re.search(r'[a-zA-Z]{3,}', stripped):
        return True
    return False

with open("english_text_pages222_242.txt", "w", encoding="utf-8") as f:
    f.write("=" * 80 + "\n")
    f.write("ENGLISH TEXT FROM APPROXIMATE PAGES 222-242\n")
    f.write("Paragraph range: 7300-8172\n")
    f.write("=" * 80 + "\n\n")
    
    # Also find any page markers in this range
    f.write("--- PAGE MARKERS IN RANGE ---\n")
    for i in range(7300, 8200):
        para = doc.paragraphs[i]
        text = para.text.strip()
        # Look for footer/header patterns like "Class - 8 : Social Science  XXX  Chapter Name"
        if re.search(r'Class\s*-?\s*8', text, re.IGNORECASE):
            f.write(f"  [{i}] '{text[:150]}'\n")
    
    f.write("\n--- ALL ENGLISH PARAGRAPHS ---\n\n")
    
    count = 0
    for i in range(7300, 8175):
        para = doc.paragraphs[i]
        text = para.text.strip()
        if is_real_english(text):
            style = para.style.name
            f.write(f"[PARA {i}] Style='{style}'\n")
            f.write(f"{text}\n")
            f.write("-" * 60 + "\n")
            count += 1
    
    f.write(f"\nTotal English paragraphs found: {count}\n")

print("Output written to english_text_pages222_242.txt")
