"""
Fix grammatical and spelling errors in English text of
'8th Social Text Book SEm1 2026.docx' (pages 222-242).
Only English paragraphs are modified. Telugu text is untouched.
All formatting, images, and styles are preserved.
"""
import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
import re
import copy
import os

def replace_in_paragraph(paragraph, old_text, new_text):
    """
    Replace old_text with new_text in a paragraph, preserving run formatting.
    Returns True if replacement was made, False otherwise.
    """
    runs = paragraph.runs
    if not runs:
        return False
    
    # Get full paragraph text from runs
    full_text = ''.join(run.text for run in runs)
    if old_text not in full_text:
        return False
    
    # Find position of old_text
    start_pos = full_text.index(old_text)
    end_pos = start_pos + len(old_text)
    
    # Map character positions to runs
    char_count = 0
    for i, run in enumerate(runs):
        run_start = char_count
        run_end = char_count + len(run.text)
        
        if run_start <= start_pos < run_end:
            if end_pos <= run_end:
                # old_text is entirely within this run — simple case
                before = run.text[:start_pos - run_start]
                after = run.text[end_pos - run_start:]
                run.text = before + new_text + after
                return True
            else:
                # old_text spans multiple runs
                # Put replacement in the first run, clear from subsequent runs
                before = run.text[:start_pos - run_start]
                run.text = before + new_text
                remaining = end_pos - run_end
                for j in range(i + 1, len(runs)):
                    if remaining <= 0:
                        break
                    rlen = len(runs[j].text)
                    if remaining >= rlen:
                        remaining -= rlen
                        runs[j].text = ''
                    else:
                        runs[j].text = runs[j].text[remaining:]
                        remaining = 0
                return True
        
        char_count = run_end
    
    return False


def apply_corrections(doc):
    """Apply all identified corrections to the document."""
    
    corrections_log = []
    
    # Define corrections as (paragraph_index, old_text, new_text, description)
    corrections = [
        # --- Chapter 4 Questions section ---
        (7372, "Lakshmibai resistance", "Lakshmibai's resistance",
         "Missing possessive apostrophe"),
        
        # --- Chapter 5: The Kakatiya Kingdom ---
        (7483, "Thousand piller temple", "Thousand Pillar Temple",
         "Spelling: 'piller' → 'Pillar'; capitalization"),
        
        (7484, "Rigveda mention this", "Rigveda mentions this",
         "Subject-verb agreement: singular subject needs 'mentions'"),
        
        (7484, "Ikshvakus Pallavas and", "Ikshvakus, Pallavas and",
         "Missing comma in list"),
        
        (7485, "There emerged a five important kingdoms emerged in south India.",
         "Five important kingdoms emerged in south India.",
         "Removed duplicate 'emerged' and incorrect article 'a'"),
        
        (7485, "The Hoyasalas", "The Hoysalas",
         "Spelling: 'Hoyasalas' → 'Hoysalas'"),
        
        (7485, "one belongs to", "one belonged to",
         "Tense: should be past tense 'belonged'"),
        
        (7485, "their own effort.", "their own efforts.",
         "Pluralization: 'effort' → 'efforts'"),
        
        (7487, "among the the Andhras", "among the Andhras",
         "Removed duplicate 'the'"),
        
        (7487, "for some extent", "to some extent",
         "Preposition: 'for some extent' → 'to some extent'"),
        
        (7488, "served as a feudatories", "served as feudatories",
         "Removed incorrect article 'a' before plural noun"),
        
        (7488, "remarkable to the telugu literature", "remarkable to Telugu literature",
         "Capitalization: 'telugu' → 'Telugu'; removed unnecessary 'the'"),
        
        (7488, "sculptures to the telugu land", "sculptures to the Telugu land",
         "Capitalization: 'telugu' → 'Telugu'"),
        
        (7535, "capital.Rudra deva", "capital. Rudra Deva",
         "Missing space after period; capitalization of 'Deva'"),
        
        (7540, "The Region of the Prola -II was a land mark",
         "The reign of Prola II was a landmark",
         "Spelling: 'Region' → 'reign'; 'land mark' → 'landmark'; fixed spacing"),
        
        (7540, "the Chalaukyas", "the Chalukyas",
         "Spelling: 'Chalaukyas' → 'Chalukyas'"),
        
        (7586, "Rudra deva defeated", "Rudra Deva defeated",
         "Capitalization: 'deva' → 'Deva'"),
        
        (7592, "Ganapathi deva was", "Ganapathi Deva was",
         "Capitalization: 'deva' → 'Deva'"),
        
        (7665, "in Telugu Sanskrit languages", "in Telugu and Sanskrit languages",
         "Missing conjunction 'and'"),
        
        (7667, "Rudrama devi Came to throne in1262 CE",
         "Rudrama Devi came to the throne in 1262 CE",
         "Capitalization fixes; lowercase 'came'; added 'the'; added space before '1262'"),
        
        (7667, "the kakatiya kingdom", "the Kakatiya kingdom",
         "Capitalization: 'kakatiya' → 'Kakatiya'"),
        
        (7725, "Marco polo", "Marco Polo",
         "Capitalization: 'polo' → 'Polo'"),
        
        (7738, "available even today also", "available even today",
         "Removed redundant 'also'"),
        
        (7739, "Other rich section also made gifts like land, others. tanks, money, cattle, jewellry etc.",
         "Other rich sections also made gifts like land, tanks, money, cattle, jewellery, etc.",
         "Plural 'sections'; removed garbled 'others.'; spelling 'jewellery'"),
        
        (7741, "in paritcular", "in particular",
         "Spelling: 'paritcular' → 'particular'"),
        
        (7741, "Vaishnavisim", "Vaishnavism",
         "Spelling: 'Vaishnavisim' → 'Vaishnavism'"),
        
        (7794, "being combinedly worshipped", "being worshipped together",
         "'combinedly' is non-standard; replaced with 'worshipped together'"),
        
        (7841, "Prataparudra- II 's", "Prataparudra II's",
         "Fixed spacing and punctuation"),
        
        (7841, "Delhi sulthans", "Delhi Sultans",
         "Spelling and capitalization: 'sulthans' → 'Sultans'"),
        
        (7841, "in 1323C.E", "in 1323 CE",
         "Added space; standard 'CE' format"),
        
        (7841, "taken prisioner", "taken prisoner",
         "Spelling: 'prisioner' → 'prisoner'"),
        
        (7842, "emerged in coastal Andhra Small kingdoms such as Kondveedu, Rajahmundry, Kandukuru etc emerged. These kindoms",
         "emerged in coastal Andhra. These kingdoms",
         "Removed duplicated sentence; fixed 'kindoms' → 'kingdoms'"),
        
        (7842, "from muslim invasions", "from Muslim invasions",
         "Capitalization: 'muslim' → 'Muslim'"),
        
        (7844, "They United by the with the stand the Muslim onslought.",
         "They united to withstand the Muslim onslaught.",
         "Fixed garbled sentence; spelling: 'onslought' → 'onslaught'"),
        
        (7844, "The Manusuri nayakas", "The Musunuri Nayakas",
         "Spelling: 'Manusuri' → 'Musunuri'; capitalization: 'nayakas' → 'Nayakas'"),
        
        (7844, "The Padmanayakas Telangana area", "The Padmanayakas ruled the Telangana area",
         "Added missing verb 'ruled' and article 'the'"),
        
        (7844, "velamakonda", "Velamakonda",
         "Capitalization"),
        
        (7845, "into telugu", "into Telugu",
         "Capitalization: 'telugu' → 'Telugu'"),
        
        (7845, "palnati Charitra", "Palnati Charitra",
         "Capitalization: 'palnati' → 'Palnati'"),
        
        (7845, "Bhama khadambham", "Bhama Khadambham",
         "Capitalization: 'khadambham' → 'Khadambham'"),
        
        (7845, "kasi khadambham", "Kasi Khadambham",
         "Capitalization: 'kasi khadambham' → 'Kasi Khadambham'"),
        
        (7845, "under the role of", "under the rule of",
         "Spelling: 'role' → 'rule'"),
        
        (7902, "Duringthe preriod", "During the period",
         "Missing space; spelling: 'preriod' → 'period'"),
        
        (7902, "were flourished", "flourished",
         "'flourish' is intransitive — removed incorrect passive 'were'"),
        
        (7903, "Bayyaram explains about Kakatiya's rule",
         "Bayyaram explain Kakatiya rule",
         "Subject-verb agreement: 'explains' → 'explain'; removed unnecessary 'about' and possessive"),
        
        (7905, "kingdons", "kingdoms",
         "Spelling: 'kingdons' → 'kingdoms'"),
        
        (7979, "does not correctly explains", "does not correctly explain",
         "Subject-verb agreement after 'does'"),
        
        (7986, "combinedly worshiped", "worshipped together",
         "'combinedly' is non-standard; fixed 'worshiped' → 'worshipped'"),
        
        (7987, "combinedly worshiped", "worshipped together",
         "'combinedly' is non-standard; fixed 'worshiped' → 'worshipped'"),
        
        (8070, "and appreciates the role", "and appreciate the role",
         "Parallel structure: 'Analyse... and appreciate' (base form)"),
        
        (8149, "of the Loksabha", "of the Lok Sabha",
         "'Lok Sabha' should be two words"),
    ]
    
    # Apply corrections
    success_count = 0
    fail_count = 0
    
    for para_idx, old_text, new_text, description in corrections:
        para = doc.paragraphs[para_idx]
        result = replace_in_paragraph(para, old_text, new_text)
        if result:
            corrections_log.append(f"✅ PARA {para_idx}: {description}")
            corrections_log.append(f"   OLD: {old_text}")
            corrections_log.append(f"   NEW: {new_text}")
            success_count += 1
        else:
            corrections_log.append(f"❌ PARA {para_idx}: FAILED - text not found")
            corrections_log.append(f"   LOOKING FOR: {old_text}")
            # Show what the paragraph actually contains
            actual = ''.join(run.text for run in para.runs)
            if not actual:
                actual = para.text
            corrections_log.append(f"   ACTUAL TEXT: {actual[:200]}")
            fail_count += 1
    
    return corrections_log, success_count, fail_count


def main():
    input_file = "8th Social Text Book SEm1 2026.docx"
    output_file = "8th Social Text Book SEm1 2026_CORRECTED.docx"
    
    print(f"Loading document: {input_file}")
    print("This may take a moment for a 53MB file...")
    doc = Document(input_file)
    print(f"Document loaded. Total paragraphs: {len(doc.paragraphs)}")
    
    print("\nApplying corrections...")
    log, success, fail = apply_corrections(doc)
    
    # Print log
    print("\n" + "=" * 70)
    print("CORRECTION LOG")
    print("=" * 70)
    for line in log:
        print(line)
    
    print("\n" + "=" * 70)
    print(f"SUMMARY: {success} corrections applied, {fail} failed")
    print("=" * 70)
    
    # Save corrected document
    print(f"\nSaving corrected document as: {output_file}")
    doc.save(output_file)
    print(f"✅ Done! File saved: {output_file}")
    
    # Save log to file
    with open("correction_log.txt", "w", encoding="utf-8") as f:
        f.write("CORRECTION LOG\n")
        f.write("=" * 70 + "\n")
        for line in log:
            f.write(line + "\n")
        f.write("\n" + "=" * 70 + "\n")
        f.write(f"SUMMARY: {success} corrections applied, {fail} failed\n")
    
    print(f"📝 Correction log saved to: correction_log.txt")


if __name__ == "__main__":
    main()
