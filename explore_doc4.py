"""
Explore document content after page 207 marker (para 7014).
Write output to file to avoid encoding issues.
"""
from docx import Document
import re

doc = Document("8th Social Text Book SEm1 2026.docx")

with open("explore_results.txt", "w", encoding="utf-8") as f:
    f.write(f"Total paragraphs: {len(doc.paragraphs)}\n\n")
    
    # Since page numbers 222-242 aren't in paragraph text,
    # they must be auto-generated in headers/footers.
    # Let's look at ALL content from para 7014 (page 207) onwards
    # to understand the structure and find the right range.
    
    # First, find ALL standalone page-like numbers in entire doc
    f.write("=== ALL STANDALONE NUMBERS (potential page markers) ===\n")
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text and re.match(r'^\d{1,3}\.?$', text):
            f.write(f"  [{i}] '{text}' | Style='{para.style.name}'\n")
    
    # Show content from para 7014 onwards (200 paras with text)
    f.write("\n=== CONTENT FROM PARA 7014 ONWARDS (200 text paras) ===\n")
    shown = 0
    for i in range(7014, min(len(doc.paragraphs), 10000)):
        para = doc.paragraphs[i]
        text = para.text.strip()
        if text and shown < 200:
            has_telugu = bool(re.search(r'[\u0C00-\u0C7F]', text))
            lang = "TE" if has_telugu else "EN"
            display = text[:150] + ("..." if len(text) > 150 else "")
            f.write(f"  [{i}] [{lang}] Style='{para.style.name}' | {display}\n")
            shown += 1
    
    # Also look further ahead
    f.write("\n=== CONTENT FROM PARA 8000 ONWARDS (200 text paras) ===\n")
    shown = 0
    for i in range(8000, min(len(doc.paragraphs), 12000)):
        para = doc.paragraphs[i]
        text = para.text.strip()
        if text and shown < 200:
            has_telugu = bool(re.search(r'[\u0C00-\u0C7F]', text))
            lang = "TE" if has_telugu else "EN"
            display = text[:150] + ("..." if len(text) > 150 else "")
            f.write(f"  [{i}] [{lang}] Style='{para.style.name}' | {display}\n")
            shown += 1

    f.write("\nDone!\n")

print("Output written to explore_results.txt")
