"""
Script to explore the structure of the textbook document.
Goal: Find paragraphs between printed page numbers 222 and 242,
and understand how English vs Telugu text is arranged.
"""
from docx import Document
import re

doc = Document("8th Social Text Book SEm1 2026.docx")

print(f"Total paragraphs: {len(doc.paragraphs)}")
print("=" * 80)

# First, let's find paragraphs that contain the printed page numbers 222-242
# Page numbers are often in headers/footers or standalone paragraphs
# Let's search for them

# Look for page number patterns in paragraphs
page_number_indices = []
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    # Look for standalone page numbers or page numbers in text
    if text and re.match(r'^\d{1,3}$', text):
        num = int(text)
        if 200 <= num <= 260:
            page_number_indices.append((i, num, text))
            print(f"Paragraph {i}: Page number = {num}, Style='{para.style.name}'")

print("\n" + "=" * 80)
print("Looking for page numbers embedded in text...")

# Also check for page numbers that might be part of larger text
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text and ('222' in text or '242' in text) and len(text) < 50:
        print(f"Paragraph {i}: '{text}' | Style='{para.style.name}'")

print("\n" + "=" * 80)
print("\nSearching for section/chapter markers near expected range...")

# Let's also look at a range of paragraphs to understand structure
# First, find approximate location by looking at all standalone numbers
all_page_nums = []
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text and re.match(r'^\d{1,3}$', text):
        num = int(text)
        if 100 <= num <= 300:
            all_page_nums.append((i, num))

print("\nAll standalone page numbers found (100-300 range):")
for idx, num in all_page_nums:
    print(f"  Paragraph index {idx}: Page {num}")

print("\n" + "=" * 80)
# Now let's look at the content around page 222
if all_page_nums:
    # Find the paragraph index for page 222
    target_indices = [(idx, num) for idx, num in all_page_nums if num == 222]
    if target_indices:
        start_idx = target_indices[0][0]
        print(f"\nContent around page 222 (paragraph {start_idx}):")
        print("-" * 60)
        for i in range(max(0, start_idx - 2), min(len(doc.paragraphs), start_idx + 15)):
            para = doc.paragraphs[i]
            text = para.text.strip()
            if text:
                # Show first 100 chars
                display = text[:100] + ("..." if len(text) > 100 else "")
                print(f"  [{i}] Style='{para.style.name}' | {display}")
